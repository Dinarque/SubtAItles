# -*- coding: utf-8 -*-
"""
Created on Tue Sep 12 17:49:25 2023

@author: 3b13j
"""

debug_mode = False
import ffmpeg
import streamlit as st 
from pytube import YouTube
from streamlit_player import st_player
import os
import shutil
import random 
import io
from docx import Document
import pysrt
from moviepy.editor import VideoFileClip, TextClip, CompositeVideoClip
from moviepy.video.tools.subtitles import SubtitlesClip


## some functions 

def get_video(link):
    # get the video and interesting metadata from the youtube link
    
    yt = YouTube(link)
    video = yt.streams.first()
    
    try :
        os.remove("buffer.mp4")
    except :
        print("clear")
    try :
         os.remove("buffer.mp3")
    except :
         print("clear")
    try :
        os.remove("final_video.mp4")
    except :
          print("clear")
    out_file = video.download()
    mp4 = shutil.copy(out_file, "buffer.mp4") 
    mp3 = shutil.copy("buffer.mp4", "buffer.mp3") 
       
    if debug_mode : 
        st.subheader("LS")
        st.write(os.listdir())
    
    return yt.title, yt.thumbnail_url, mp3, mp4
   
   


def get_subtitles(link):
    #Use Assemblyai to generate the subtitles for the video. We LOVE it ! 
    import assemblyai as aai
    try :
        aai.settings.api_key = st.secrets["KEY"]
    except : 
        st.write("no more funds ")
    transcriber = aai.Transcriber()
    transcript = transcriber.transcribe(link)
    return transcript
 
# the next step is to replace part of the words in the subtitle with blanks for the students to fill in? Here the difficulty referes to the percentage of deleted files
   

def treat_vtt(vtt, dif) : 
    # erases some words  in a file that is in the vtt format.
    
    
    buffer = vtt.split("\n")
    if debug_mode : 
        for i in range (len(buffer)) : 
            st.write(i, buffer[i])
    
    to_modif = {}
    for i in range(len(buffer)) :
        if i > 1 and i%3 == 0 :
            to_modif[i] = buffer[i]
    
    
    for k in to_modif.keys() :        
        line = to_modif[k].split(" ")
        n_l = ""
        for wd in line : 
           rand = random.randint(0, 101)
           if rand <= dif or  dif == 100: 
               n_l +=  len(wd) * "_"+ " "
           else : 
               n_l += wd + " "    
        buffer[k] = n_l
    
    if debug_mode : 
        for i in range (len(buffer)) : 
            st.write(i, buffer[i])
   
    truffer = ""
    for line in buffer : truffer += line +"\n"
    return truffer



def treat_srt(srt, dif) : 
    
    # erases some words  in a file that is in the srt format.
    buffer = srt.split("\n")
    if debug_mode : 
        for i in range (len(buffer)) : 
            st.write(i, buffer[i])
    
    to_modif = {}
    for i in range(len(buffer)) :
        if  (i+2)%4 == 0 :
            to_modif[i] = buffer[i]
            
            

    for k in to_modif.keys() : 
        
        line = to_modif[k].split(" ")
        n_l = ""
        for wd in line : 
           rand = random.randint(0, 101)
           if rand < dif or  dif == 100 : 
               n_l +=  len(wd) * "_"+ " "
           else : 
               n_l += wd + " "
           
        buffer[k] = n_l
    
    if debug_mode : 
        for i in range (len(buffer)) : 
            st.write(i, buffer[i])
   
    truffer = ""
    for i in range(int(len(buffer)/4)) :
       truffer += buffer[4*i] +" "+ buffer[4*i+1] +" "+ buffer[4*i+2]+" "+ buffer[4*i+3]+ "\n"
    
    return truffer



def subtitles_to_file(obj, form, dif) :
    # global function that converts the subtitles into a fill-in-the-blanks activity in the chosen format
    if form == "vtt" : res = treat_vtt(obj.export_subtitles_vtt(), dif)
    elif form == "srt" : res = treat_srt(obj.export_subtitles_srt(), dif)
    else : st.write("invalid subtitles format")
    
    with open("subtitles."+str(form), "w") as file:
        file.write(res)
    return res



def add_subtitles(mp4, sub) : 
    # function (in progress) to incorporate the subtitles into the video.
    video = ffmpeg.input(mp4)
    audio = video.audio
    ffmpeg.concat(video.filter("subtitles", sub), audio, v=1, a=1).output("final_video.mp4").run()
    return "final_video.mp4"
    """
    (
    ffmpeg
    .input(mp4)
    .filter('subtitles', sub)
    .output('final_vodep.mp4')
    .run()
    )
    st.write("caribou polu")
    """
    return 'final_video.mp4'
    
def save_subtitles_to_doc(subs) : 
    #transform the subtitles into a worksheet that the user can download
    import docx
    doc = docx.Document()
    doc.add_heading(f'Worksheet for the video "{st.session_state.title}"',1)
    doc.add_paragraph(subs)
    doc.save('worksheet.docx')
    
    
    

"""
attemps to merge subtitles and video

def time_to_seconds(time_obj):
    return time_obj.hours * 3600 + time_obj.minutes * 60 + time_obj.seconds + time_obj.milliseconds / 1000


def create_subtitle_clips(subtitles, videosize,fontsize=24, font='Arial', color='yellow', debug = False):
    subtitle_clips = []

    for subtitle in subtitles:
        start_time = time_to_seconds(subtitle.start)
        end_time = time_to_seconds(subtitle.end)
        duration = end_time - start_time

        video_width, video_height = videosize
        
        text_clip = TextClip(subtitle.text, fontsize=fontsize, font=font, color=color, bg_color = 'black',size=(video_width*3/4, None), method='caption').set_start(start_time).set_duration(duration)
        subtitle_x_position = 'center'
        subtitle_y_position = video_height* 4 / 5 

        text_position = (subtitle_x_position, subtitle_y_position)                    
        subtitle_clips.append(text_clip.set_position(text_position))

    return subtitle_clips

"""

    
## sidebar 

if debug_mode : 
    st.session_state 

with st.sidebar :
    st.header("Configurate the app")
    "___"
    
    link = st.text_input("Put link to the youtube video :")
    st.write("Choose the % of words that will be hidden (you can modify it even once the video is loaded)")
    difficulty = st.slider('Adjust difficulty', 0, 100)
    st.session_state.difficulty = difficulty
    
    if st.button("Lets learn !") :
        #main part that runs the analysis 
        st.session_state.link = link
        st.session_state.title, st.session_state.thumbnail, st.session_state.mp3 , st.session_state.mp4 = get_video(link)
        st.write("file ready")
        st.session_state.subtitles = get_subtitles(st.session_state.mp3)
        
        

if "link" in st.session_state : 
    st_player(link)
"___"
if "subtitles" in st.session_state : 
    st.subheader("subtitles :") 
    processed_sub = subtitles_to_file(st.session_state.subtitles, "srt", st.session_state.difficulty)
    st.write(processed_sub)
    save_subtitles_to_doc(processed_sub)
   
    doc = Document('worksheet.docx')
    doc_download = doc
    bio = io.BytesIO()
    doc_download.save(bio)
    
    st.download_button(label="Download worksheet", data = bio ,
            file_name= 'worksheet.docx',
            mime="docx"
        )
    
    st.write("lets play the video")
    
    video_file = open(st.session_state.mp4, 'rb')
    video_bytes = video_file.read()
    st.video(video_bytes)
    st.write("lets add the subtitles")
    #final = add_subtitles(st.session_state.mp4, "subtitles.srt")
    #st.video(final)
    
    
    generator = lambda txt: TextClip(txt, font='Arial', fontsize=16, color='white')
    subtitles = SubtitlesClip("subtitles.srt", generator)

    video = VideoFileClip(st.session_state.mp4)
    result = CompositeVideoClip([video, subtitles.set_pos(('center','bottom'))])

    result.write_videofile("out.mp4", fps=video.fps, temp_audiofile="temp-audio.m4a", remove_temp=True, codec="libx264", audio_codec="aac")


    """
    Attemps to merge subtitles and video
    
    video = VideoFileClip(st.session_state.mp4)
    subtitles = pysrt.open("subtitles.srt")
    
    
    output_video_file = "final_video.mp4"
    
    # Create subtitle clips
    subtitle_clips = create_subtitle_clips(subtitles,video.size)
    
    # Add subtitles to the video
    final_video = CompositeVideoClip([video] + subtitle_clips)
    
    # Write output video file
    final_video.write_videofile(output_video_file)
        
    st.video("final_video.mp4")
    """  
        
else : 
    st.subheader("Welcome to SubtAItle")
    '___'
    "This app is designed to turn any Youtube video into a language learning opportunity : It uses AssemblyAI to generate subtitles for the video, then erases part of the subtitles to create a fill-in-the blanks activity."
    "The activity can then be downloaded and filled out while watching the video. The goal would also have been to merge the subtitles with the video but the humble developper I am didn't find the right way to proceed"
    "___"
    st.image("")
    "WARNING : it is not allowed to use all the videos available on Youtube since some of them are under copyright protection. It is your job as an responsible user to ensure you only use the App with free of right content "
    " For instance you can use all videos under 'fair use' and all the videos that have a 'download' button on their page"
    "___"
    "This app was built using the AssemblyAI technlogy."
    "It cost around 1 cent for every minute of video transcribed. See https://www.assemblyai.com/pricing for more information on pricing"
    "
    st.subheader("Please update a video")
    
        
        
    




    
    